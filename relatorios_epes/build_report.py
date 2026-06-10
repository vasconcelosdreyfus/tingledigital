#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report.py — Pipeline de geração dos relatórios mensais Tingle Digital → EPES.

Uso:
    python3 build_report.py <input.md> <output.docx> [--reference ref.docx]
    python3 build_report.py --all          # gera os 6 relatórios em /tmp/epes_relatorios/final

Etapas:
    1. Pré-processa o markdown (remove headings de título duplicados na capa,
       remove o bloco de assinatura textual, promove níveis de heading quando
       o corpo não tem H1).
    2. pandoc md -> docx.
    3. Pós-processa com python-docx: página A4 + margens, styles (Heading 1/2/3,
       corpo, legendas), tabelas (header framboesa, zebra, sem bordas verticais),
       capa (banner EPES + título + faixa multicolor + logo Tingle no rodapé da
       1ª página), header/footer institucionais com logo e nº de página, caixa
       de aviso ("INSERIR FOTOS"), e bloco de assinatura padronizado com imagem.
"""

import argparse
import copy
import os
import re
import subprocess
import sys
import tempfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------- constantes

BASE = "/tmp/epes_relatorios"
ASSETS = os.path.join(BASE, "design", "assets")
BANNER = os.path.join(ASSETS, "banner_epes.jpg")
LOGO_TINGLE = os.path.join(ASSETS, "logo_tingle.png")
ASSINATURA = os.path.join(ASSETS, "assinatura.png")

# Paleta Tingle
RASPBERRY = "C8385E"
AMBER = "F5A623"
GREEN = "7AB648"
TEAL = "1E9E96"
TEAL_DARK = "17807A"
BROWN = "4A2E24"
NEON = "58E05C"
BODY_GREY = "333333"
GREY_MID = "808080"
GREY_LIGHT = "D9D9D9"
ZEBRA = "F5F5F7"

PAGE_W, PAGE_H = Cm(21.0), Cm(29.7)
MARGIN_LR = Cm(2.2)
MARGIN_TOP = Cm(3.0)
MARGIN_BOTTOM = Cm(2.3)
CONTENT_W = Cm(21.0 - 2 * 2.2)  # 16.6 cm

FOOTER_TEXT = ("TINGLEDIGITAL SERVICOS LTDA – RUA VISC DE PIRAJA 414, SAL 718 "
               "– IPANEMA – RIO DE JANEIRO – CEP: 22410-905")

SIGNATURE_DATE = "Rio de Janeiro, 9 de junho de 2026."
SIGNATURE_LINES = [
    ("Dreyfus Vasconcelos", True),
    ("Sócio Administrador", False),
    ("Tingledigital Serviços Ltda", False),
    ("CNPJ 33.486.049/0001-55", False),
    ("Rua Visconde De Piraja, 414 SL 718 — IPANEMA — RIO DE JANEIRO RJ", False),
    ("CEP 22.410-002", False),
]

# Configuração por relatório
REPORTS = {
    "001_maio.md": dict(
        title="Relatório Técnico da Consultoria para a\nPrestação de Contas da Casa Brasil",
        subtitle="Termo de Colaboração nº 003/2023 · 1º Aditivo · Processo Administrativo nº 11.363/2023",
        header="Prestação de Contas · Maio 2026",
        strip_headings=2, promote=True,
        out="202605 Relatório Técnico Prestação de Contas.docx",
    ),
    "002_maio.md": dict(
        title="Monitoramento, Avaliação e\nSustentabilidade da Proposta",
        subtitle="Termo de Colaboração nº 003/2023 · 1º Aditivo · Processo Administrativo nº 11.363/2023",
        header="Monitoramento e Avaliação · Maio 2026",
        strip_headings=1, promote=False,
        out="2605 RELATÓRIO DE MONITORAMENTO.docx",
    ),
    "003_maio.md": dict(
        title="Workshop InovaTalks EPES 2026",
        subtitle="Relatório Técnico do Ciclo de Palestras · Termo de Colaboração nº 003/2023 · 1º Aditivo",
        header="Workshop InovaTalks · Maio 2026",
        strip_headings=3, promote=True,
        out="202605 Workshop InovaTalks EPES 2026.docx",
    ),
    "004_maio.md": dict(
        title="Relatório Técnico da Participação\nno RIO2C 2026",
        subtitle="Termo de Colaboração nº 003/2023 · 1º Aditivo · Cidade das Artes — Rio de Janeiro",
        header="RIO2C 2026 · Maio 2026",
        strip_headings=1, promote=True,
        out="202605 Evento RIO2C.docx",
    ),
    "005_maio.md": dict(
        title="Relatório Técnico do\nServiço de Sistema EAD",
        subtitle="Plataforma Cognita · Termo de Colaboração nº 003/2023 · 1º Aditivo",
        header="Sistema EAD · Maio 2026",
        strip_headings=2, promote=True,
        out="202605 Relatório Serviço de Sistema EAD.docx",
    ),
    "006_maio.md": dict(
        title="Relatório Técnico de Desenvolvimento\nde Material Pedagógico",
        subtitle="Termo de Colaboração nº 003/2023 · 1º Aditivo · Processo Administrativo nº 11.363/2023",
        header="Material Pedagógico · Maio 2026",
        strip_headings=1, promote=True,
        out="202605 Relatório Material Pedagogico para a EPES.docx",
    ),
}


# ------------------------------------------------------------- xml helpers

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def set_shading(parent_pr, fill):
    for old in parent_pr.findall(qn("w:shd")):
        parent_pr.remove(old)
    parent_pr.append(_el("w:shd", val="clear", color="auto", fill=fill))


def set_para_border(paragraph, edges, color, sz=4, space=1, val="single"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for edge in edges:
        e = _el("w:" + edge, val=val, sz=sz, space=space, color=color)
        pBdr.append(e)


def set_cell_borders(tc, spec):
    """spec: dict edge -> (val, sz, color) ; val 'nil' to remove."""
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, (val, sz, color) in spec.items():
        for old in borders.findall(qn("w:" + edge)):
            borders.remove(old)
        if val == "nil":
            borders.append(_el("w:" + edge, val="nil"))
        else:
            borders.append(_el("w:" + edge, val=val, sz=sz, color=color))


def add_page_field(paragraph, size_pt, color):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = _el("w:sz", val=int(size_pt * 2))
    col = _el("w:color", val=color)
    rPr.append(col)
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def has_drawing(paragraph):
    return paragraph._p.find(".//" + qn("w:drawing")) is not None


# ------------------------------------------------------- markdown pre-proc

def preprocess_markdown(md_path, cfg):
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()

    # 1. remove o bloco de assinatura textual (da última linha "Rio de Janeiro, ..." em diante)
    idx = text.rfind("\nRio de Janeiro,")
    if idx != -1:
        text = text[:idx].rstrip() + "\n"

    # 2. remove os headings de título do topo (irão para a capa)
    lines = text.split("\n")
    removed, out = 0, []
    for ln in lines:
        if removed < cfg["strip_headings"] and re.match(r"^#{1,2}\s", ln):
            removed += 1
            continue
        if removed < cfg["strip_headings"] and ln.strip() == "":
            continue  # pula linhas em branco entre headings de título
        out.append(ln)
    text = "\n".join(out)

    # 3. promove níveis de heading se o corpo não tem H1
    if cfg["promote"]:
        text = re.sub(r"^(#{2,6})\s", lambda m: "#" * (len(m.group(1)) - 1) + " ", text, flags=re.M)

    return text


# ------------------------------------------------------------ docx styling

def style_base(doc):
    st = doc.styles
    normal = st["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BODY_GREY)
    pf = normal.paragraph_format
    pf.line_spacing = 1.2
    pf.space_after = Pt(6)

    for name in ("Body Text", "First Paragraph"):
        try:
            s = st[name]
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            s.paragraph_format.space_after = Pt(7)
        except KeyError:
            pass

    try:
        s = st["Compact"]
        s.paragraph_format.space_after = Pt(3)
    except KeyError:
        pass

    specs = [
        ("Heading 1", 17, RASPBERRY, Pt(22), Pt(8)),
        ("Heading 2", 13.5, TEAL_DARK, Pt(15), Pt(6)),
        ("Heading 3", 11.5, BROWN, Pt(11), Pt(4)),
        ("Heading 4", 10.5, GREY_MID, Pt(9), Pt(3)),
    ]
    for name, size, color, before, after in specs:
        try:
            h = st[name]
        except KeyError:
            continue
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.italic = False
        h.font.color.rgb = RGBColor.from_string(color)
        h.paragraph_format.space_before = before
        h.paragraph_format.space_after = after
        h.paragraph_format.keep_with_next = True

    # linha âmbar fina sob cada Heading 1 (definida no style)
    try:
        h1 = st["Heading 1"]
        pPr = h1.element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pBdr.append(_el("w:bottom", val="single", sz=6, space=3, color=AMBER))
        pPr.append(pBdr)
    except KeyError:
        pass

    for name in ("Image Caption", "Caption", "Table Caption"):
        try:
            s = st[name]
        except KeyError:
            continue
        s.font.size = Pt(9)
        s.font.italic = True
        s.font.color.rgb = RGBColor.from_string(GREY_MID)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(16)
        s.paragraph_format.space_before = Pt(4)

    try:
        s = st["Captioned Figure"]
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.space_before = Pt(10)
    except KeyError:
        pass


def setup_page(doc):
    for sec in doc.sections:
        sec.page_width = PAGE_W
        sec.page_height = PAGE_H
        sec.left_margin = MARGIN_LR
        sec.right_margin = MARGIN_LR
        sec.top_margin = MARGIN_TOP
        sec.bottom_margin = MARGIN_BOTTOM
        sec.header_distance = Cm(1.1)
        sec.footer_distance = Cm(1.0)


NUM_RE = re.compile(r"^[\s\d.,%+–—-]*\d[\s\d.,%]*$|^R\$\s*[\d.,]+$")


def style_tables(doc):
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr
        # largura 100%
        for old in tblPr.findall(qn("w:tblW")):
            tblPr.remove(old)
        tblPr.append(_el("w:tblW", w=5000, type="pct"))
        # margens internas das células
        for old in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(old)
        mar = OxmlElement("w:tblCellMar")
        for side, w in (("top", 50), ("bottom", 50), ("left", 110), ("right", 110)):
            mar.append(_el("w:" + side, w=w, type="dxa"))
        tblPr.append(mar)

        rows = table.rows
        ncols = len(rows[0].cells) if rows else 0

        # larguras proporcionais ao conteúdo (corrige colunas mal distribuídas do pandoc)
        if ncols:
            weights = []
            for j in range(ncols):
                maxlen = max((len(row.cells[j].text.strip())
                              for row in rows if j < len(row.cells)), default=4)
                weights.append(min(max(maxlen, 6), 40))
            total_tw = Cm(16.6).twips
            widths = [int(total_tw * w / sum(weights)) for w in weights]
            grid = tbl.find(qn("w:tblGrid"))
            if grid is not None:
                for gc in list(grid):
                    grid.remove(gc)
                for w in widths:
                    grid.append(_el("w:gridCol", w=w))
            for row in rows:
                for j, cell in enumerate(row.cells):
                    if j < len(widths):
                        tcPr = cell._tc.get_or_add_tcPr()
                        for old in tcPr.findall(qn("w:tcW")):
                            tcPr.remove(old)
                        tcPr.append(_el("w:tcW", w=widths[j], type="dxa"))

        # detecta colunas numéricas
        numeric_cols = []
        for j in range(ncols):
            vals = []
            for row in rows[1:]:
                if j < len(row.cells):
                    vals.append(row.cells[j].text.strip())
            non_empty = [v for v in vals if v]
            if non_empty and all(NUM_RE.match(v.replace("**", "")) for v in non_empty):
                numeric_cols.append(j)

        for i, row in enumerate(rows):
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(_el("w:cantSplit"))  # linha nunca quebra entre páginas
            if i == 0:
                trPr.append(_el("w:tblHeader", val="true"))
            for j, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                edge_spec = {
                    "left": ("nil", 0, ""), "right": ("nil", 0, ""),
                    "top": ("single", 4, GREY_LIGHT),
                    "bottom": ("single", 4, GREY_LIGHT),
                }
                if i == 0:
                    edge_spec["top"] = ("single", 8, RASPBERRY)
                    edge_spec["bottom"] = ("single", 8, RASPBERRY)
                    set_shading(tcPr, RASPBERRY)
                elif i == len(rows) - 1:
                    edge_spec["bottom"] = ("single", 8, "A6A6A6")
                if i > 0 and i % 2 == 0:
                    set_shading(tcPr, ZEBRA)
                set_cell_borders(cell._tc, edge_spec)

                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.line_spacing = 1.1
                    if j in numeric_cols:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.size = Pt(9.5)
                        if i == 0:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor.from_string("FFFFFF")


def style_figures(doc):
    """Centraliza qualquer parágrafo que contenha imagem inline."""
    for p in doc.paragraphs:
        if has_drawing(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(10)


def style_notice_boxes(doc):
    for p in doc.paragraphs:
        if "[INSERIR FOTOS" in p.text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            set_para_border(p, ("top", "bottom", "left", "right"), AMBER, sz=6, space=8)
            set_shading(p._p.get_or_add_pPr(), "FDF3E0")
            for r in p.runs:
                r.font.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor.from_string("8A6116")


# --------------------------------------------------------------- signature

def add_signature(doc):
    # remove parágrafos vazios sobrando no fim (inclui &nbsp; herdados dos md)
    while doc.paragraphs and not doc.paragraphs[-1].text.strip() \
            and not has_drawing(doc.paragraphs[-1]):
        doc.paragraphs[-1]._p.getparent().remove(doc.paragraphs[-1]._p)

    # encadeia o último parágrafo de texto do corpo ao bloco de assinatura
    # (w:keepNext): ou a assinatura cabe na última página de conteúdo, ou o
    # parágrafo final da conclusão desce junto — nunca assinatura órfã
    blocks = [c for c in doc.element.body
              if c.tag in (qn("w:p"), qn("w:tbl"))]
    if blocks and blocks[-1].tag == qn("w:p"):
        pPr = blocks[-1].get_or_add_pPr()
        if pPr.find(qn("w:keepNext")) is None:
            pPr.append(_el("w:keepNext"))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(SIGNATURE_DATE)
    r.font.size = Pt(10.5)

    pimg = doc.add_paragraph()
    pimg.paragraph_format.space_before = Pt(8)
    pimg.paragraph_format.space_after = Pt(0)
    pimg.paragraph_format.keep_with_next = True
    pimg.add_run().add_picture(ASSINATURA, width=Cm(5.0))

    # linha fina sob a assinatura
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(6)
    rule.paragraph_format.keep_with_next = True
    rule.add_run(" ").font.size = Pt(2)
    set_para_border(rule, ("top",), GREY_LIGHT, sz=4, space=1)
    rule.paragraph_format.right_indent = Cm(9.6)  # linha de ~7cm

    for i, (line, bold) in enumerate(SIGNATURE_LINES):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.line_spacing = 1.15
        if i < len(SIGNATURE_LINES) - 1:
            sp.paragraph_format.keep_with_next = True
        r = sp.add_run(line)
        r.font.size = Pt(10.5) if i < 2 else Pt(9)
        r.font.bold = bold
        if bold:
            r.font.color.rgb = RGBColor.from_string(BROWN)
        elif i >= 2:
            r.font.color.rgb = RGBColor.from_string("595959")


# ----------------------------------------------------------- header/footer

def build_header_footer(doc, cfg):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True

    # ---- header páginas internas
    hdr = sec.header
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(LOGO_TINGLE, width=Cm(1.6))
    rt = p.add_run("\t" + cfg["header"])
    rt.font.size = Pt(9)
    rt.font.color.rgb = RGBColor.from_string(GREY_MID)
    set_para_border(p, ("bottom",), GREY_LIGHT, sz=4, space=4)

    # ---- footer páginas internas
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    for r in list(fp.runs):
        r._r.getparent().remove(r._r)
    fp.paragraph_format.tab_stops.clear_all()
    fp.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
    set_para_border(fp, ("top",), GREY_LIGHT, sz=4, space=4)
    fr = fp.add_run(FOOTER_TEXT + "\t")
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor.from_string(GREY_MID)
    add_page_field(fp, 8, GREY_MID)

    # ---- 1ª página: header vazio, footer = logo + tagline (rodapé da capa)
    fph = sec.first_page_header
    fph.is_linked_to_previous = False
    for r in list(fph.paragraphs[0].runs):
        r._r.getparent().remove(r._r)

    fpf = sec.first_page_footer
    fpf.is_linked_to_previous = False
    p1 = fpf.paragraphs[0]
    for r in list(p1.runs):
        r._r.getparent().remove(r._r)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run().add_picture(LOGO_TINGLE, width=Cm(2.8))
    p2 = fpf.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("TINGLE DIGITAL · RELATÓRIO TÉCNICO MENSAL")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(GREY_MID)
    r2.font.bold = True


# -------------------------------------------------------------------- cover

def build_cover(doc, cfg):
    """Insere a capa antes do primeiro elemento do corpo."""
    body = doc.element.body
    anchor = body[0]
    created = []

    def para():
        p = doc.add_paragraph()
        created.append(p._p)
        return p

    # banner full-width
    pb = para()
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(4)
    pb.add_run().add_picture(BANNER, width=CONTENT_W)
    set_para_border(pb, ("bottom",), NEON, sz=14, space=6)

    # espaçador
    sp = para()
    sp.paragraph_format.space_before = Pt(70)
    sp.paragraph_format.space_after = Pt(0)
    sp.add_run(" ").font.size = Pt(2)

    # título
    for chunk in cfg["title"].split("\n"):
        pt = para()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(2)
        pt.paragraph_format.line_spacing = 1.05
        rt = pt.add_run(chunk)
        rt.font.size = Pt(24)
        rt.font.bold = True
        rt.font.name = "Calibri"
        rt.font.color.rgb = RGBColor.from_string(BROWN)

    # faixa multicolor (tabela 4 células)
    tbl = doc.add_table(rows=1, cols=4)
    created.append(tbl._tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblPr.append(_el("w:tblW", w=0, type="auto"))
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(_el("w:trHeight", val=90, hRule="exact"))
    for cell, color in zip(tbl.rows[0].cells, (RASPBERRY, AMBER, GREEN, TEAL)):
        cell.width = Cm(2.1)
        tcPr = cell._tc.get_or_add_tcPr()
        set_shading(tcPr, color)
        set_cell_borders(cell._tc, {e: ("nil", 0, "") for e in
                                    ("top", "bottom", "left", "right")})
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.space_before = Pt(0)
        cr = cp.add_run(" ")
        cr.font.size = Pt(1)

    # subtítulo
    ps = para()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(16)
    ps.paragraph_format.space_after = Pt(4)
    rs = ps.add_run(cfg["subtitle"])
    rs.font.size = Pt(10.5)
    rs.font.color.rgb = RGBColor.from_string("595959")

    pm = para()
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm.paragraph_format.space_before = Pt(6)
    rm = pm.add_run("MAIO DE 2026")
    rm.font.size = Pt(15)
    rm.font.bold = True
    rm.font.color.rgb = RGBColor.from_string(TEAL_DARK)

    # quebra de página
    pbreak = para()
    pbreak.paragraph_format.space_after = Pt(0)
    pbreak.add_run().add_break(WD_BREAK.PAGE)

    for el in created:
        anchor.addprevious(el)


# --------------------------------------------------------------------- main

def build(md_path, out_path, reference=None):
    name = os.path.basename(md_path)
    cfg = REPORTS.get(name)
    if cfg is None:
        sys.exit(f"Sem configuração para {name} (edite REPORTS em build_report.py)")

    text = preprocess_markdown(md_path, cfg)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_md = os.path.join(tmp, "in.md")
        tmp_docx = os.path.join(tmp, "out.docx")
        with open(tmp_md, "w", encoding="utf-8") as f:
            f.write(text)
        cmd = ["pandoc", "-f", "markdown+smart", tmp_md, "-o", tmp_docx]
        if reference:
            cmd += ["--reference-doc", reference]
        subprocess.run(cmd, check=True)

        doc = Document(tmp_docx)
        setup_page(doc)
        style_base(doc)
        style_tables(doc)
        style_figures(doc)
        style_notice_boxes(doc)
        add_signature(doc)
        build_header_footer(doc, cfg)
        build_cover(doc, cfg)

        doc.core_properties.title = cfg["title"].replace("\n", " ")
        doc.core_properties.author = "Tingle Digital"
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
        doc.save(out_path)
    print(f"OK  {out_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md", nargs="?", help="markdown de entrada")
    ap.add_argument("out", nargs="?", help="docx de saída")
    ap.add_argument("--reference", help="reference.docx opcional para o pandoc")
    ap.add_argument("--all", action="store_true", help="gera os 6 relatórios em /tmp/epes_relatorios/final")
    args = ap.parse_args()

    if args.all:
        final = os.path.join(BASE, "final")
        for md_name, cfg in REPORTS.items():
            build(os.path.join(BASE, "out", md_name),
                  os.path.join(final, cfg["out"]), args.reference)
    else:
        if not args.md or not args.out:
            ap.error("informe <md> <out.docx> ou --all")
        build(args.md, args.out, args.reference)


if __name__ == "__main__":
    main()
